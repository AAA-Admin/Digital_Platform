from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
BLACK  = RGBColor(0x1A, 0x1A, 0x2E)   # near-black heading
ACCENT = RGBColor(0xC9, 0x9A, 0x06)   # gold accent
MUTED  = RGBColor(0x55, 0x55, 0x66)   # muted body

def set_run_font(run, size, bold=False, italic=False, color=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(text, level=1, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_run_font(run, sizes.get(level, 11), bold=True, color=color)
    return p

def add_body(text, bold=False, italic=False, color=MUTED, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, 10, bold=bold, italic=italic, color=color)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_run_font(r1, 10, bold=True, color=BLACK)
        r2 = p.add_run(text)
        set_run_font(r2, 10, color=MUTED)
    else:
        run = p.add_run(text)
        set_run_font(run, 10, color=MUTED)

def shade_row(row, hex_color="F5F0E8"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = t.rows[0]
    shade_row(hdr, "1A1A2E")
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run  = p.add_run(h)
        set_run_font(run, 9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = "FFFFFF" if ri % 2 == 0 else "F9F6F0"
        shade_row(row, fill)
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run  = p.add_run(str(cell_text))
            set_run_font(run, 9, color=BLACK)

    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Cm(width)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "C99A06")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_callout(text):
    """Gold-bordered note box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(8)
    run = p.add_run("  " + text)
    set_run_font(run, 10, italic=True, color=RGBColor(0x6B, 0x4C, 0x00))
    # left border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    lft  = OxmlElement("w:left")
    lft.set(qn("w:val"),   "single")
    lft.set(qn("w:sz"),    "18")
    lft.set(qn("w:space"), "8")
    lft.set(qn("w:color"), "C99A06")
    pBdr.append(lft)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
#  TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("AAA EVENTS & PRODUCTION")
set_run_font(r, 20, bold=True, color=BLACK)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("Project Agreement")
set_run_font(r2, 13, color=ACCENT)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(2)
r3 = p3.add_run("Prepared by Core Tensor  ·  March 2026")
set_run_font(r3, 9, italic=True, color=MUTED)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 1 — WHAT WE ARE DOING
# ══════════════════════════════════════════════════════════════════════════════
add_heading("PART 1 — WHAT WE ARE DOING", level=1, color=ACCENT)
add_body(
    "Everything in this section is confirmed, in scope, and included in the agreed ₹28,000. "
    "No extras, no surprises.",
    bold=True, color=BLACK
)

# ── Section 1: Confirmed Scope ────────────────────────────────────────────────
add_heading("1. Confirmed Scope", level=2)
add_table(
    ["#", "Deliverable", "Detail"],
    [
        ["1", "Premium Immersive Website",       "Custom-coded, high-performance landing page"],
        ["2", "Basic SEO & Google Indexing",      "On-page SEO so you appear on Page 1 for your target keywords"],
        ["3", "Lead Capture Form + WhatsApp Alert","Visitors fill a form; all 3 owners get an instant WhatsApp message with lead details"],
    ],
    col_widths=[0.7, 4.5, 7.8]
)

add_body("Agreed Investment: ₹28,000", bold=True, color=BLACK)
add_callout(
    "Important: SEO & Google Indexing and the Lead Capture Form with WhatsApp alerts are "
    "not add-ons. Both are fully built and delivered within the ₹28,000. There are no extra "
    "charges for either."
)

# ── Section 2: Why No Templates ───────────────────────────────────────────────
add_heading("2. Why We Do Not Use Templates", level=2)
add_body(
    "Platforms like Wix, Squarespace, or WordPress load 50–200 unnecessary scripts. "
    "We refuse to use them for AAA Events for four reasons:",
    color=MUTED
)
add_bullet("Speed — templates load in 3–6 seconds. We guarantee under 0.8 seconds.", bold_prefix="Speed.")
add_bullet("SEO ceiling — bloated pages are structurally penalised by Google's Core Web Vitals.", bold_prefix="SEO ceiling.")
add_bullet("Brand authority — large-scale event clients lose trust the moment they see a Wix-style page.", bold_prefix="Brand authority.")
add_bullet("Ownership — Wix/Squarespace means renting. We build on infrastructure you own forever.", bold_prefix="Ownership.")
doc.add_paragraph()

# ── Section 3: Performance Guarantee ─────────────────────────────────────────
add_heading("3. Performance Guarantee", level=2)
add_callout(
    "We guarantee your website will load in under 0.8 seconds on a standard 4G mobile connection. "
    "Demonstrated live on Google PageSpeed Insights and GTmetrix before handover. "
    "If we miss this, we fix it at zero cost."
)
add_body("Achieved by:", color=MUTED)
add_bullet("Hand-crafted HTML/CSS/JS — zero bloat")
add_bullet("Global CDN (Vercel) — your site is served from a node near your visitor")
add_bullet("Compressed and lazy-loaded images")
add_bullet("No unnecessary third-party scripts")
doc.add_paragraph()

# ── Section 4: Lead Capture Form ─────────────────────────────────────────────
add_heading("4. Lead Capture Form", level=2)
add_body("Visitors on your website fill in:", color=MUTED)
add_table(
    ["Field", "Type"],
    [
        ["Name",                  "Text input"],
        ["Phone Number",          "Text input"],
        ["City / Location",       "Dropdown"],
        ["Type of Event",         "Dropdown (Concert / Corporate / Wedding / Exhibition / Other)"],
        ["Approximate Budget Range", "Dropdown"],
        ["Brief Requirement",     "Short text area"],
    ],
    col_widths=[5.5, 7.5]
)
add_body("On submission, all three owners instantly receive a WhatsApp message:", color=MUTED)

# WhatsApp message sample
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.8)
p.paragraph_format.space_after  = Pt(8)
sample = (
    "🔔 NEW LEAD — AAA Events Website\n\n"
    "Name: Rahul Sharma\n"
    "Phone: +91 98765 43210\n"
    "City: Bengaluru\n"
    "Event Type: Concert\n"
    "Budget: ₹5L – ₹15L\n"
    "Requirement: Need LED wall + truss setup for 2,000 pax...\n\n"
    "Reply directly to follow up."
)
run = p.add_run(sample)
set_run_font(run, 9, color=RGBColor(0x1A, 0x1A, 0x2E))
run.font.name = "Courier New"

# ── Section 5: Accounts & Infrastructure ──────────────────────────────────────
add_heading("5. Accounts & Infrastructure — Yours, Not Ours", level=2)
add_body(
    "Everything is set up under your name. We guide the setup. You own it forever. "
    "We never sit between you and any service.",
    color=MUTED
)
add_table(
    ["Service", "What It's For", "Account Owner", "Cost"],
    [
        ["Domain (GoDaddy / Namecheap)", "Your web address (e.g. aaaevents.in)", "AAA Events", "~₹800–1,500/year"],
        ["Vercel (Hosting)",             "Serves your site globally",             "AAA Events", "Free"],
        ["WhatsApp Business API (Meta)", "Sends form alerts to your WhatsApp",    "AAA Events", "Free (1,000 msgs/month)"],
        ["Google Search Console",        "Tracks your Google rankings",           "AAA Events", "Free"],
        ["Google Analytics",             "Tracks visitor behaviour",              "AAA Events", "Free"],
    ],
    col_widths=[3.5, 4.5, 3.0, 2.0]
)
add_body(
    "Monthly running cost: ~₹0 for the foreseeable future. Only cost is domain renewal (~₹1,200/year).",
    bold=True, color=BLACK
)

# ── Section 6: SEO Coverage ───────────────────────────────────────────────────
add_heading("6. SEO — What Is Included at ₹28,000", level=2)
add_body("All of the following is included. No extra charge.", color=MUTED)

seo_in = [
    "Google Search Console setup — site verified and submitted",
    "XML Sitemap generated and submitted",
    "On-Page SEO — title tags, meta descriptions, heading structure, image alt tags",
    'Target keywords: "event production company Bengaluru", "LED wall rental Bengaluru", "concert stage setup company", and up to 8 total',
    "Schema Markup (LocalBusiness) — helps with Google Map Pack",
    "Core Web Vitals compliance — passing speed and usability tests",
    "robots.txt — correct crawl instructions",
    "Open Graph tags — so your link looks good on WhatsApp and social",
]
for item in seo_in:
    add_bullet(item)
doc.add_paragraph()

# ── Section 7: Full In/Out Scope ──────────────────────────────────────────────
add_heading("7. Full Scope — What Is In & What Is Out", level=2)
add_body("✅  IN SCOPE — Included in ₹28,000", bold=True, color=RGBColor(0x1A, 0x7A, 0x2E))

in_scope = [
    ("Website", [
        "Single-page custom-coded landing page (HTML/CSS/JS — no templates, no WordPress)",
        "Mobile-first responsive design",
        "Scroll-triggered animations and reveal effects",
        "AAA logo animation on page load",
        "Services / offerings section",
        "Founders / team section",
        "Google Maps embed",
        "Gallery / portfolio section (using images you provide)",
        "Contact details section",
        "Deployed to Vercel on your account",
    ]),
    ("Lead Capture Form", [
        "One embedded form (fields as described above)",
        "Instant WhatsApp alert to all 3 owners on submission",
        "Basic spam protection (honeypot field)",
        "Confirmation message shown to user on submit",
    ]),
    ("SEO & Indexing", [
        "Google Search Console setup and verification",
        "XML sitemap and robots.txt",
        "Title tags, meta descriptions, image alt tags",
        "LocalBusiness schema markup",
        "Open Graph tags",
        "Target keyword research (up to 8 keywords)",
        "Core Web Vitals audit and pass before handover",
    ]),
    ("Project Management", [
        "One round of design revisions after first draft",
        "One round of copy/content revisions",
        "Handover call + walkthrough of all accounts",
        "30-day post-launch bug support (bugs only, not new features)",
    ]),
]
for category, items in in_scope:
    add_body(category, bold=True, color=BLACK, space_after=2)
    for item in items:
        add_bullet(item)
doc.add_paragraph()

add_body("❌  OUT OF SCOPE — Not Included, Requires Separate Agreement", bold=True, color=RGBColor(0xCC, 0x22, 0x22))
add_table(
    ["What", "Why It's Separate"],
    [
        ["Google My Business setup",          "Best done after site is live — separate engagement"],
        ["Backlink building",                  "Ongoing work — separate engagement"],
        ["Paid ads (Google / Meta)",           "Requires ad budget + management fee"],
        ["Blog or content writing",            "Recurring monthly work — separate engagement"],
        ["Second page or additional pages",    "Single-page scope only"],
        ["Booking / payment gateway",          "Separate engagement"],
        ["Custom admin panel",                 "Separate engagement"],
        ["WhatsApp Bot (automated reply flow)","Separate engagement"],
        ["AI lead research / database",        "Separate engagement"],
        ["Social media management",            "Not a web development service"],
        ["Logo / brand identity work",         "Separate unless agreed otherwise"],
        ["Video or photography",               "Content provided by you"],
        ["Email marketing setup",              "Not scoped"],
        ["Monthly maintenance retainer",       "Quoted separately if required"],
    ],
    col_widths=[6.0, 7.0]
)

add_body("Revision policy:", bold=True, color=BLACK, space_after=2)
add_bullet("Up to 2 rounds of design revisions, 1 round of content/copy revisions during build")
add_bullet("After handover: 30-day bug support only. New features or changed requirements = new agreement.")
doc.add_paragraph()

# ── Section 8: Payment Terms ──────────────────────────────────────────────────
add_heading("8. Payment Terms", level=2)
add_table(
    ["Milestone", "Amount", "When"],
    [
        ["Advance",     "₹8,000",  "On agreement confirmation, before work begins"],
        ["Mid-project", "₹10,000", "On first draft delivery — design live for your review"],
        ["Final",       "₹10,000", "After every IN SCOPE item is delivered and you sign off"],
        ["TOTAL",       "₹28,000", ""],
    ],
    col_widths=[4.0, 3.0, 6.0]
)
add_bullet("No payment is due before the previous milestone is delivered")
add_bullet("Final payment released only after you have reviewed and approved all deliverables")
add_bullet("No automatic charges — every future engagement is a separate agreement")

# ── Section 9: Summary ────────────────────────────────────────────────────────
add_heading("9. Summary", level=2)
add_table(
    ["Item", "Detail"],
    [
        ["Scope",                  "Custom website + Basic SEO + Form with WhatsApp alerts"],
        ["Agreed Price",           "₹28,000"],
        ["Payment Schedule",       "₹8,000 advance → ₹10,000 on first draft → ₹10,000 on sign-off"],
        ["Monthly Running Cost",   "~₹0 (only domain renewal ~₹1,200/year)"],
        ["All Third-Party Accounts","In your name, billed to you — we never sit in between"],
        ["Page Load Guarantee",    "Under 0.8 seconds — demonstrated before handover"],
        ["No Templates",           "100% custom code, owned by you"],
        ["Timeline",               "10–14 working days from content receipt"],
    ],
    col_widths=[5.0, 8.0]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 2 — WHAT ELSE WE CAN DO (FUTURE)
# ══════════════════════════════════════════════════════════════════════════════
add_heading("PART 2 — WHAT ELSE WE CAN DO", level=1, color=ACCENT)
add_body(
    "None of this is in the current ₹28,000 engagement. These are future capabilities — "
    "available when the business is ready. Each is a separate conversation and a separate agreement.",
    italic=True, color=MUTED
)

# ── Traffic Expectations ───────────────────────────────────────────────────────
add_heading("A. What to Realistically Expect from Traffic", level=2)
add_body(
    "A website alone does not generate traffic — it converts traffic. SEO is your signboard on Google. "
    "Here is an honest timeline:",
    color=MUTED
)
add_table(
    ["Timeframe", "What Happens", "Why"],
    [
        ["Week 1–2",    "Google discovers and indexes the site",                        "We submit it directly; Google crawls within days"],
        ["Month 1–3",   "You appear in searches, likely page 2–5",                      "New domain has zero history — Google is still building trust"],
        ["Month 4–6",   "Page 1 rankings for low-competition keywords",                 'e.g. "truss setup company Bengaluru" — specific, less fought-over'],
        ["Month 6–12",  "Consistent page 1 for 3–5 keywords",                          "Realistic ceiling with on-page SEO only"],
        ["Month 12+",   "Compounding growth",                                           "Every backlink, review, or content piece accelerates ranking"],
    ],
    col_widths=[2.5, 5.0, 5.5]
)
add_body(
    "Honest traffic estimate at Month 6 (on-page SEO only): 150–400 visitors/month. "
    "That is 10–25 genuine inbound leads/month if the site converts well — which ours will.",
    italic=True, color=MUTED
)
add_callout(
    "Free action you can take right now: Share your website link across every Instagram post, "
    "WhatsApp group, personal profile, and client conversation. This builds early traffic signals that help SEO at zero cost."
)

# ── GMB ────────────────────────────────────────────────────────────────────────
add_heading("B. Google My Business (GMB)", level=2)
add_body(
    "GMB is the box that appears on the right side of Google and in Google Maps. "
    "For a local B2B company like AAA Events, GMB is often more valuable than the website itself for local discovery. "
    'When someone searches "event company near me" or "stage setup Bengaluru", the Map Pack of 3 businesses is the first thing they see.',
    color=MUTED
)
add_body("What a standalone GMB engagement covers:", bold=True, color=BLACK, space_after=2)
add_bullet("Verifying your business address with Google (postcard or video call)")
add_bullet("Uploading 20+ photos and writing optimised business descriptions")
add_bullet("Setting up your first review-gathering process (5+ reviews in Month 1 = strong early signal)")
add_body(
    "Best done after the website is live — so your GMB links to a credible, fast-loading page.",
    italic=True, color=MUTED
)
doc.add_paragraph()

# ── Backlinks ──────────────────────────────────────────────────────────────────
add_heading("C. Backlink Building", level=2)
add_body(
    "A backlink is when another website links to yours. Google treats each one as a vote of confidence. "
    "A new business typically gets 0–5 organic backlinks in Year 1 — that's normal.",
    color=MUTED
)
add_body("What a backlink engagement covers:", bold=True, color=BLACK, space_after=2)
add_bullet("Submission to 15–20 legitimate Indian event & business directories (each = a free backlink)")
add_bullet("A press-style case study of your best event, optimised to attract natural links")
add_bullet("Identifying partnership sites (venues, caterers, AV suppliers) who can link to you")
add_callout(
    "Free action: Ask your top 5 clients to leave a Google Review on your GMB page once it's live. "
    "Five genuine 5-star reviews outperform most paid SEO activities."
)

# ── Growth Levers ──────────────────────────────────────────────────────────────
add_heading("D. Other Growth Levers Available", level=2)
add_table(
    ["Growth Lever", "Impact", "Notes"],
    [
        ["Google My Business",      "Highest for local B2B — Map Pack visibility",   "Separate engagement — best after site launch"],
        ["Backlink Campaign",        "Tells Google you're credible",                  "Separate engagement"],
        ["Content / Blog",           "Each article is a new entry point into your site","Recurring monthly work"],
        ["Paid Ads (Google/Meta)",   "Instant traffic, stops when budget stops",      "Your ad budget + our management fee"],
        ["WhatsApp Bot",             "Automated reply flow to incoming leads",        "Separate engagement"],
        ["Booking / Payment Gateway","Let clients book and pay online",               "Separate engagement"],
        ["Admin Panel",              "Edit your own website without a developer",     "Separate engagement"],
        ["AI Lead Research",         "Identify and qualify leads automatically",      "Separate engagement"],
        ["Monthly SEO Retainer",     "Ongoing ranking improvement",                   "Quoted separately"],
    ],
    col_widths=[4.0, 5.0, 4.0]
)

# ── Scalability ────────────────────────────────────────────────────────────────
add_heading("E. Built to Grow — Scalability Without a Rebuild", level=2)
add_body(
    "The website is built on a modern serverless architecture designed to expand. "
    "When the business is ready, new capabilities — automated lead flows, booking systems, admin portals, "
    "AI integrations — can be layered on top of the existing foundation without rebuilding anything.",
    color=MUTED
)
add_body(
    "Nothing gets rebuilt from scratch. Every rupee invested now carries forward.",
    bold=True, color=BLACK
)

add_divider()

# ── Footer ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Core Tensor builds custom, performance-first digital systems. "
    "We do not resell hosting, take platform kickbacks, or lock you into subscriptions. "
    "Everything we build is yours."
)
set_run_font(r, 8, italic=True, color=MUTED)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/home/user/AAA/01_Client_Facing/AAA_Events_Project_Agreement.docx"
doc.save(out)
print(f"Saved: {out}")
