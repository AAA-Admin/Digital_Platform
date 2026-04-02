from docx import Document

filepath = r"c:\Users\yuvar\OneDrive\Documents\GitHub\AAA\01_Client_Facing\AAA_Events_Project_Agreement.docx"
doc = Document(filepath)

def replace_text_in_runs(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        # Paragraphs are made of runs, if standard replace is too complex across runs,
        # we can just wipe the paragraph and set it to the new text to preserve some formatting
        # but wiping runs loses bolding. A safer way is to replace text in runs directly.
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
        
        # If the text spans multiple runs, the above won't catch it. 
        # A simpler approach: if old_text is in paragraph.text, clear runs and add back the modified text.
        # This loses inline styling for the replaced part, but keeps the paragraph style.
        if old_text in paragraph.text and not any(old_text in r.text for r in paragraph.runs):
            full_text = paragraph.text.replace(old_text, new_text)
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = full_text
                else:
                    run.text = ""

# For simplicity, we define all replacements:
replacements = {
    "all 3 owners": "both directors",
    "all three owners": "both directors",
    "scaffolding setup": "event infrastructure setup",
    "scaffolding": "event infrastructure",
    "₹8,000 advance → ₹10,000 on first draft → ₹10,000 on sign-off": "₹5,000 advance (received) → ₹13,000 on first draft → ₹10,000 on sign-off",
    "On agreement confirmation, before work begins": "Received on the agreement date; work commences on this date"
}

for para in doc.paragraphs:
    for old, new in replacements.items():
        if old in para.text:
            replace_text_in_runs(para, old, new)

for table in doc.tables:
    for row in table.rows:
        # Check if it's the specific payment row
        row_texts = [c.text for c in row.cells]
        is_mid_project_row = "Mid-project" in row_texts
        is_advance_row = "Advance" in row_texts
        
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in replacements.items():
                    if old in para.text:
                        replace_text_in_runs(para, old, new)
                
                # Manual fixes for payment amounts in the table
                if is_mid_project_row and "₹10,000" in para.text:
                    replace_text_in_runs(para, "₹10,000", "₹13,000")
                if is_advance_row:
                    if "₹8,000" in para.text:
                        replace_text_in_runs(para, "₹8,000", "₹5,000")
                    if "Advance" == para.text.strip():
                        replace_text_in_runs(para, "Advance", "Advance Received")

doc.save(filepath)
print("Updated DOCX successfully.")
