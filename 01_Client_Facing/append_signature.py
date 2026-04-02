import os
from docx import Document

filepath = r"c:\Users\yuvar\OneDrive\Documents\GitHub\AAA\01_Client_Facing\AAA_Events_Project_Agreement.docx"
doc = Document(filepath)

doc.add_page_break()
doc.add_heading('Signatures', level=2)

p1 = doc.add_paragraph("This agreement is executed in two (2) original copies, one for the Client and one for Core Tensor. By signing below, both parties confirm their understanding and acceptance of the scope, payment terms, and conditions outlined in this document.")

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run("For AAA Events & Production:").bold = True
doc.add_paragraph()
doc.add_paragraph()
p3 = doc.add_paragraph("_________________________________________\n")
p3.add_run("Authorized Signatory (Director)\n").bold = True
p3.add_run("Date: _______________")

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.add_run("For Core Tensor:").bold = True
doc.add_paragraph()
doc.add_paragraph()
p5 = doc.add_paragraph("_________________________________________\n")
p5.add_run("Authorized Signatory\n").bold = True
p5.add_run("Date: _______________")

doc.save(filepath)
print("Appended signature section to DOCX successfully.")
